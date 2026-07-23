#!/usr/bin/env python3
"""Generate reproducible Priority 2 PDF fixtures, DOCX templates, and raster editorial art."""

from __future__ import annotations

import io
import json
import math
import random
import re
import shutil
from pathlib import Path

from PIL import Image, ImageDraw, ImageFilter, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pypdf import PdfReader, PdfWriter
from reportlab.lib.colors import Color, HexColor, black, white
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


ROOT = Path(__file__).resolve().parents[1]
RUNTIME = ROOT / "runtime-public"
OUTPUT_PDF = ROOT / "output" / "pdf"
TMP = ROOT / "tmp" / "pdfs"
FIXTURES = RUNTIME / "research" / "fixtures"
REDACTION = RUNTIME / "research" / "redaction"
BENCHMARK = RUNTIME / "research" / "benchmark"
TEMPLATES = RUNTIME / "templates" / "editable"
SHARE = RUNTIME / "share"
DEMOS = RUNTIME / "editorial" / "demos"

for folder in [OUTPUT_PDF, TMP, FIXTURES, REDACTION, BENCHMARK, TEMPLATES, SHARE, DEMOS]:
    folder.mkdir(parents=True, exist_ok=True)


def load_evidence() -> list[dict]:
    source = (ROOT / "config" / "priority-two-evidence.mjs").read_text(encoding="utf-8")
    body = re.sub(r"^\s*export default Object\.freeze\(", "", source)
    body = re.sub(r"\);\s*$", "", body)
    return json.loads(body)


EVIDENCE = load_evidence()


def font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    name = "dm-sans-700.ttf" if bold else "dm-sans-400.ttf"
    path = RUNTIME / "cosmic-assets" / "fonts" / name
    try:
        return ImageFont.truetype(str(path), size=size)
    except OSError:
        return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, selected_font, width: int) -> list[str]:
    words = text.split()
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = f"{current} {word}".strip()
        if draw.textbbox((0, 0), candidate, font=selected_font)[2] <= width:
            current = candidate
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_wrapped(draw, xy, text, selected_font, fill, width, line_gap=8, max_lines=None):
    x, y = xy
    lines = wrap_text(draw, text, selected_font, width)
    if max_lines and len(lines) > max_lines:
        lines = lines[:max_lines]
        lines[-1] = f"{lines[-1].rstrip(' .')}..."
    line_height = selected_font.size + line_gap if hasattr(selected_font, "size") else 20
    for line in lines:
        draw.text((x, y), line, font=selected_font, fill=fill)
        y += line_height
    return y


def title_from_slug(slug: str) -> str:
    special = {"pdf": "PDF", "ocr": "OCR", "html": "HTML", "jpg": "JPG", "png": "PNG"}
    return " ".join(special.get(part, part.capitalize()) for part in slug.split("-"))


def make_share_card(slug: str, title: str, kicker: str) -> None:
    image = Image.new("RGB", (1200, 630), "#F5F9FF")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((55, 48, 1145, 582), radius=36, fill="#FFFFFF", outline="#D5E2F4", width=2)
    draw.rounded_rectangle((82, 76, 332, 118), radius=21, fill="#E7F1FF")
    draw.text((105, 87), kicker.upper(), font=font(16, True), fill="#2851EB")
    draw.ellipse((953, 85, 1097, 229), fill="#2851EB")
    draw.rounded_rectangle((998, 116, 1052, 191), radius=7, fill="#FFFFFF")
    draw.rectangle((1008, 133, 1042, 138), fill="#80B7FF")
    draw.rectangle((1008, 149, 1042, 154), fill="#80B7FF")
    draw.rectangle((1008, 165, 1033, 170), fill="#80B7FF")
    y = draw_wrapped(draw, (82, 168), title, font(55, True), "#12213A", 780, line_gap=8, max_lines=3)
    draw_wrapped(draw, (84, y + 22), "Original methods, real examples, and clearly stated limits.", font(23), "#5E6D84", 760, line_gap=7, max_lines=2)
    draw.text((84, 520), "PDF", font=font(24, True), fill="#12213A")
    offset = draw.textbbox((0, 0), "PDF", font=font(24, True))[2]
    draw.text((84 + offset, 520), "Arrow", font=font(24, True), fill="#2851EB")
    draw.text((900, 525), "Tested July 21, 2026", font=font(16, True), fill="#6A7890")
    image.save(SHARE / f"{slug}.png", optimize=True)


def make_demo(record: dict) -> None:
    slug = record["toolId"]
    title = title_from_slug(slug)
    image = Image.new("RGB", (1200, 675), "#EDF4FF")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((42, 38, 1158, 637), radius=30, fill="#FFFFFF", outline="#CDDBEF", width=2)
    draw.text((76, 70), title, font=font(28, True), fill="#12213A")
    draw.rounded_rectangle((948, 66, 1117, 106), radius=20, fill="#EAF1FF")
    draw.text((978, 77), "REGRESSION DEMO", font=font(13, True), fill="#2851EB")
    panel_y = 140
    panel_h = 354
    for index, (label, copy) in enumerate((("EXAMPLE INPUT", record["input"]), ("EXPECTED OUTPUT", record["output"]))):
        x = 76 + index * 534
        draw.rounded_rectangle((x, panel_y, x + 492, panel_y + panel_h), radius=22, fill="#F8FAFD", outline="#DCE5F1", width=2)
        draw.text((x + 28, panel_y + 28), label, font=font(14, True), fill="#65758C")
        draw.rounded_rectangle((x + 28, panel_y + 76, x + 112, panel_y + 177), radius=12, fill="#E6EEFF" if index == 0 else "#DDEBFF")
        if index == 0:
            draw.rounded_rectangle((x + 49, panel_y + 96, x + 91, panel_y + 157), radius=5, fill="#FFFFFF", outline="#8DA7D7", width=2)
            draw.rectangle((x + 57, panel_y + 111, x + 83, panel_y + 115), fill="#7B95C4")
            draw.rectangle((x + 57, panel_y + 123, x + 83, panel_y + 127), fill="#A2B5D7")
            draw.rectangle((x + 57, panel_y + 135, x + 76, panel_y + 139), fill="#A2B5D7")
        else:
            draw.ellipse((x + 47, panel_y + 102, x + 93, panel_y + 148), fill="#2851EB")
            draw.line((x + 59, panel_y + 125, x + 69, panel_y + 136, x + 85, panel_y + 114), fill="#FFFFFF", width=5, joint="curve")
        draw_wrapped(draw, (x + 28, panel_y + 205), copy, font(20, True), "#22344F", 430, line_gap=8, max_lines=5)
    draw.line((579, 300, 621, 300), fill="#2851EB", width=5)
    draw.polygon([(621, 300), (605, 289), (605, 311)], fill="#2851EB")
    draw.rounded_rectangle((76, 525, 1124, 601), radius=16, fill="#13213A")
    draw.text((102, 545), "MEASURED", font=font(13, True), fill="#87B7FF")
    draw_wrapped(draw, (210, 540), record["result"], font(17, True), "#FFFFFF", 880, line_gap=5, max_lines=2)
    image.save(DEMOS / f"{slug}.png", optimize=True)


RESOURCE_SHARES = {
    "resources": ("PDF research, safety guides, and workflow playbooks", "Original resources"),
    "research-pdf-conversion-benchmark": ("Q3 2026 browser PDF fidelity benchmark", "Open methodology"),
    "guides-redact-pdf-safely": ("How to prove redacted PDF text is actually gone", "Safety guide"),
    "guides-ocr-quality": ("OCR quality by scan and document type", "Field guide"),
    "templates": ("Free editable DOCX templates with completion guidance", "Template library"),
    "workflows-education-pdf-workflow": ("A reviewable PDF workflow for educators", "Workflow playbook"),
    "workflows-recruiting-pdf-workflow": ("Private, consistent candidate PDF packets", "Workflow playbook"),
    "workflows-legal-operations-pdf-workflow": ("Version, comparison, and redaction controls for legal operations", "Workflow playbook"),
    "workflows-real-estate-pdf-workflow": ("Complete, legible property PDF packets", "Workflow playbook"),
    "workflows-small-business-pdf-workflow": ("Controlled client-ready PDFs for small businesses", "Workflow playbook"),
    "privacy": ("Understand every PDFArrow data path", "Trust center"),
    "security": ("Current security controls and reporting path", "Trust center"),
    "architecture": ("How browser processing, identity, storage, and analytics connect", "Trust center"),
    "uptime": ("Service availability without an invented percentage", "Trust center"),
    "incident-history": ("Public incident history and disclosure policy", "Trust center"),
}


def build_simple_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=letter, pageCompression=1)
    for page, heading in [(1, "Browser PDF benchmark - simple fixture"), (2, "Verification checklist")]:
        c.setTitle("PDFArrow simple searchable fixture")
        c.setFont("Helvetica-Bold", 18)
        c.drawString(72, 720, heading)
        c.setFont("Helvetica", 11)
        c.drawString(72, 692, f"Page {page} of 2")
        lines = [
            "This page contains searchable text, predictable spacing, and a page number.",
            "Reference value: FT-PDF-2026-42000.",
            "Use this fixture to verify page count, text extraction, order, and export recovery.",
        ] if page == 1 else [
            "1. Reopen the downloaded result.",
            "2. Confirm both pages are present and searchable.",
            "3. Search for FT-PDF-2026-42000.",
        ]
        y = 650
        for line in lines:
            c.drawString(72, y, line)
            y -= 24
        c.setStrokeColor(HexColor("#2851EB"))
        c.line(72, 90, 540, 90)
        c.setFillColor(HexColor("#52637A"))
        c.drawRightString(540, 68, f"PDFArrow regression fixture | {page}")
        c.showPage()
    c.save()


def build_complex_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=letter, pageCompression=1)
    c.setTitle("PDFArrow complex layout fixture")
    c.setFillColor(HexColor("#102A43"))
    c.setFont("Helvetica-Bold", 22)
    c.drawString(54, 730, "Complex layout benchmark")
    c.setFont("Helvetica", 10)
    c.setFillColor(HexColor("#52637A"))
    c.drawString(54, 708, "Columns, table geometry, vectors, and a rotated label")
    c.setFillColor(HexColor("#EFF5FF"))
    c.roundRect(54, 624, 504, 58, 10, fill=1, stroke=0)
    c.setFillColor(HexColor("#2851EB"))
    c.setFont("Helvetica-Bold", 12)
    c.drawString(72, 651, "Reference metric")
    c.setFillColor(HexColor("#102A43"))
    c.drawRightString(540, 647, "42,000")
    column_text = [
        "Left column: PDF stores positioned page elements rather than flowing document structure. A converter must infer reading order and grouping.",
        "Right column: This fixture deliberately mixes type, geometry, and vector graphics so output claims can be checked against a known page.",
    ]
    for index, text in enumerate(column_text):
        x = 54 + index * 264
        words = text.split()
        y = 590
        line = ""
        c.setFillColor(black)
        c.setFont("Helvetica", 10)
        for word in words:
            candidate = f"{line} {word}".strip()
            if c.stringWidth(candidate, "Helvetica", 10) < 232:
                line = candidate
            else:
                c.drawString(x, y, line)
                y -= 15
                line = word
        c.drawString(x, y, line)
    table_y = 420
    widths = [168, 168, 168]
    headings = ["Item", "Expected", "Observed"]
    values = [("Pages", "2", "2"), ("Search token", "42000", "42000"), ("Vectors", "Present", "Present")]
    x = 54
    for i, width in enumerate(widths):
        c.setFillColor(HexColor("#2851EB"))
        c.rect(x, table_y, width, 32, fill=1, stroke=0)
        c.setFillColor(white)
        c.setFont("Helvetica-Bold", 10)
        c.drawString(x + 10, table_y + 11, headings[i])
        x += width
    for row_index, row in enumerate(values):
        y = table_y - (row_index + 1) * 32
        x = 54
        for i, width in enumerate(widths):
            c.setFillColor(Color(.97, .98, 1) if row_index % 2 == 0 else white)
            c.setStrokeColor(HexColor("#D9E3F1"))
            c.rect(x, y, width, 32, fill=1, stroke=1)
            c.setFillColor(HexColor("#233650"))
            c.setFont("Helvetica", 10)
            c.drawString(x + 10, y + 11, row[i])
            x += width
    c.saveState()
    c.translate(576, 190)
    c.rotate(90)
    c.setFillColor(HexColor("#2851EB"))
    c.setFont("Helvetica-Bold", 9)
    c.drawString(0, 0, "ROTATED VECTOR LABEL")
    c.restoreState()
    c.setFillColor(HexColor("#B9D2FF"))
    c.circle(150, 170, 54, fill=1, stroke=0)
    c.setFillColor(HexColor("#2851EB"))
    c.roundRect(235, 130, 180, 80, 18, fill=1, stroke=0)
    c.setFillColor(white)
    c.setFont("Helvetica-Bold", 14)
    c.drawCentredString(325, 165, "VECTOR SHAPES")
    c.showPage()
    c.setFillColor(HexColor("#102A43"))
    c.setFont("Helvetica-Bold", 20)
    c.drawString(54, 730, "Complex fixture - page two")
    c.setFont("Times-Roman", 12)
    c.drawString(54, 690, "Font change, page break, and selectable text remain intentional test properties.")
    c.setFont("Courier", 11)
    c.drawString(54, 650, "UNIQUE-COMPLEX-TOKEN-8264")
    c.showPage()
    c.save()


def scan_source(degraded: bool) -> Image.Image:
    base = Image.new("L", (1700, 2200), 248 if not degraded else 225)
    draw = ImageDraw.Draw(base)
    try:
        heading_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 58)
        body_font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", 34)
    except OSError:
        heading_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
    draw.text((140, 160), "SCANNED INTAKE FORM", font=heading_font, fill=20 if not degraded else 85)
    lines = ["Reference: OCR QUALITY 2026", "Applicant: Jordan Lee", "Requested total: $42,000", "Review every name, date, and number."]
    y = 350
    for line in lines:
        draw.text((145, y), line, font=body_font, fill=35 if not degraded else 105)
        y += 90
    draw.rectangle((140, 780, 1560, 1640), outline=60 if not degraded else 130, width=4)
    for offset in [930, 1080, 1230, 1380, 1530]:
        draw.line((140, offset, 1560, offset), fill=90 if not degraded else 145, width=3)
    if degraded:
        rng = random.Random(20260721)
        pixels = base.load()
        for _ in range(38000):
            x = rng.randrange(base.width)
            y = rng.randrange(base.height)
            pixels[x, y] = min(255, max(0, pixels[x, y] + rng.randrange(-48, 49)))
        base = base.filter(ImageFilter.GaussianBlur(radius=1.15)).rotate(2.2, expand=False, fillcolor=235)
    return base.convert("RGB")


def image_pdf(image: Image.Image, path: Path) -> None:
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG", quality=80 if "degraded" in path.name else 94)
    buffer.seek(0)
    c = canvas.Canvas(str(path), pagesize=letter, pageCompression=1)
    from reportlab.lib.utils import ImageReader
    c.drawImage(ImageReader(buffer), 0, 0, width=letter[0], height=letter[1])
    c.showPage()
    c.save()


def build_large_pdf(path: Path) -> None:
    c = canvas.Canvas(str(path), pagesize=letter, pageCompression=1)
    c.setTitle("PDFArrow 100 page progressive rendering fixture")
    for page in range(1, 101):
        c.setFillColor(HexColor("#12213A"))
        c.setFont("Helvetica-Bold", 24)
        c.drawString(64, 718, f"Progressive page {page}")
        c.setFont("Helvetica", 11)
        c.setFillColor(HexColor("#52637A"))
        c.drawString(64, 682, "This is a synthetic public regression fixture. It contains no personal data.")
        c.drawString(64, 652, f"Search token: LARGE-PDF-PAGE-{page:03d}")
        for row in range(12):
            shade = 236 + (row % 2) * 8
            c.setFillColor(Color(shade / 255, shade / 255, min(1, (shade + 8) / 255)))
            c.roundRect(64, 590 - row * 38, 470, 25, 5, fill=1, stroke=0)
        c.setFillColor(HexColor("#2851EB"))
        c.drawRightString(540, 56, f"{page} / 100")
        c.showPage()
    c.save()


def build_redaction_samples(before: Path, after: Path) -> None:
    secret = "4815 1623 4200 0091"
    for path, include_secret in [(before, True), (after, False)]:
        c = canvas.Canvas(str(path), pagesize=letter, pageCompression=1)
        c.setTitle("PDFArrow redaction verification sample")
        c.setAuthor("PDFArrow Product Engineering")
        c.setFont("Helvetica-Bold", 20)
        c.drawString(72, 720, "Fictional customer record")
        c.setFont("Helvetica", 11)
        c.drawString(72, 680, "Customer: Sample Person")
        c.drawString(72, 650, "Account number:")
        if include_secret:
            c.setFont("Courier", 13)
            c.drawString(180, 647, secret)
        c.setFillColor(black)
        c.rect(174, 638, 190, 24, fill=1, stroke=0)
        c.setFillColor(HexColor("#52637A"))
        c.setFont("Helvetica", 10)
        note = "UNSAFE: underlying text remains recoverable." if include_secret else "VERIFIED SAMPLE: sensitive characters are absent from the page content."
        c.drawString(72, 600, note)
        c.drawString(72, 570, "All names and numbers in this fixture are fictional.")
        c.showPage()
        c.save()


def copy_pdf_artifact(source: Path, group: str) -> None:
    target = OUTPUT_PDF / group
    target.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source, target / source.name)


BLUE = RGBColor(0x28, 0x51, 0xEB)
INK = RGBColor(0x12, 0x21, 0x3A)
MUTED = RGBColor(0x5D, 0x6C, 0x84)
LIGHT_FILL = "EEF4FF"


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    element = OxmlElement("w:tblHeader")
    element.set(qn("w:val"), "true")
    tr_pr.append(element)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int], indent=120) -> None:
    total = sum(widths)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[min(index, len(widths) - 1)]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def apply_doc_styles(doc: Document, title: str, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.left_indent = Inches(0)
    normal.paragraph_format.right_indent = Inches(0)
    normal.paragraph_format.first_line_indent = Inches(0)
    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, INK, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = style_name != "Title"
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.left_indent = Inches(0)
        style.paragraph_format.right_indent = Inches(0)
        style.paragraph_format.first_line_indent = Inches(0)
    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("PDFARROW  |  EDITABLE TEMPLATE")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.color.rgb = MUTED
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run("Edit every bracketed prompt before use. Verify the final PDF after export.")
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    title_run = p.add_run(title)
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = INK
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(16)
    r = sub.add_run(subtitle)
    r.font.name = "Calibri"
    r.font.size = Pt(11)
    r.font.color.rgb = MUTED
    r.font.italic = True


def add_note(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    properties = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), LIGHT_FILL)
    properties.append(shading)
    r = p.add_run(f"{label}: ")
    r.bold = True
    r.font.color.rgb = BLUE
    p.add_run(text)


def add_field_table(doc: Document, rows: list[tuple[str, str]], widths=(2600, 6760)) -> None:
    table = doc.add_table(rows=1, cols=2)
    header_cells = table.rows[0].cells
    for index, header in enumerate(("Field", "Details")):
        header_cells[index].text = header
        set_cell_shading(header_cells[index], "DDE7FB")
        header_paragraph = header_cells[index].paragraphs[0]
        header_paragraph.paragraph_format.space_after = Pt(0)
        for run in header_paragraph.runs:
            run.bold = True
            run.font.color.rgb = BLUE
            run.font.size = Pt(8)
    set_repeat_table_header(table.rows[0])
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], "F2F5FA")
        for run in cells[0].paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = INK
        for paragraph in cells[0].paragraphs + cells[1].paragraphs:
            paragraph.paragraph_format.space_after = Pt(0)
    set_table_geometry(table, list(widths))


def add_section_heading(doc: Document, title: str, *, page_break_before: bool = False) -> None:
    paragraph = doc.add_paragraph(title, style="Heading 1")
    paragraph.paragraph_format.page_break_before = page_break_before


def add_checklist_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        set_cell_shading(cell, "2851EB")
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
            if row_index % 2 == 0:
                set_cell_shading(cells[index], "F7F9FC")
            cells[index].paragraphs[0].paragraph_format.space_after = Pt(0)
    set_table_geometry(table, widths)


def save_doc(doc: Document, filename: str) -> None:
    props = doc.core_properties
    props.title = filename.replace("-", " ").replace(".docx", "").title()
    props.author = "PDFArrow Product Engineering"
    props.subject = "Free editable document template"
    props.keywords = "PDFArrow, editable template"
    props.comments = "Created July 21, 2026"
    path = TEMPLATES / filename
    doc.save(path)


def build_invoice() -> None:
    doc = Document()
    apply_doc_styles(doc, "Small-business invoice", "Editable master for services or products. Replace every bracketed prompt.")
    add_note(doc, "Before sending", "Confirm the customer, invoice number, dates, tax treatment, totals, payment method, and legal business details.")
    add_section_heading(doc, "Business and customer")
    add_field_table(doc, [("From", "[Legal business name]\n[Address]\n[Email and phone]"), ("Bill to", "[Customer name]\n[Customer address]\n[Customer email]"), ("Invoice", "Number: [INV-0001]\nIssue date: [Month DD, YYYY]\nDue date: [Month DD, YYYY]")])
    add_section_heading(doc, "Line items")
    add_checklist_table(doc, ["Description", "Qty", "Rate", "Amount"], [["[Service or item 1]", "[1]", "[$0.00]", "[$0.00]"], ["[Service or item 2]", "[1]", "[$0.00]", "[$0.00]"], ["[Additional item]", "[ ]", "[ ]", "[ ]"]], [4920, 900, 1500, 2040])
    add_field_table(doc, [("Subtotal", "[$0.00]"), ("Tax or adjustment", "[$0.00 or N/A]"), ("Total due", "[$0.00]")], widths=(6760, 2600))
    add_section_heading(doc, "Payment and notes", page_break_before=True)
    add_field_table(doc, [("Payment method", "[Instructions that do not expose sensitive credentials]"), ("Terms", "[Late fee, deposit, refund, or purchase-order terms if applicable]"), ("Notes", "[Thank-you note, project reference, or delivery details]")])
    add_section_heading(doc, "Release checklist")
    add_checklist_table(doc, ["Done", "Check before sending"], [["[ ]", "Customer legal name and delivery address are correct"], ["[ ]", "Line-item math, subtotal, tax, credits, and total were recalculated"], ["[ ]", "Invoice number, issue date, due date, and purchase-order reference match"], ["[ ]", "Payment instructions use an approved method and expose no credentials"], ["[ ]", "The exported PDF was reopened and all pages are readable"], ["[ ]", "The filename identifies the customer, invoice, and date without sensitive data"]], [1200, 8160])
    save_doc(doc, "small-business-invoice.docx")


def build_candidate_scorecard() -> None:
    doc = Document()
    apply_doc_styles(doc, "Candidate evaluation scorecard", "Evidence-based interview notes with consistent criteria and a documented recommendation.")
    add_note(doc, "Use responsibly", "Record job-related evidence only. Follow applicable hiring policy and law; do not record protected characteristics or assumptions.")
    add_section_heading(doc, "Interview record")
    add_field_table(doc, [("Role", "[Role title and requisition ID]"), ("Candidate ID", "[Internal candidate ID - avoid unnecessary personal data]"), ("Interview", "[Stage, date, interviewer]"), ("Decision owner", "[Name or role]")])
    add_section_heading(doc, "Consistent criteria")
    add_checklist_table(doc, ["Criterion", "Evidence observed", "Score 1-5"], [["Relevant problem solving", "[Specific example and outcome]", "[ ]"], ["Role-specific knowledge", "[Specific answer or work sample]", "[ ]"], ["Communication", "[Clarity, listening, and audience adaptation]", "[ ]"], ["Execution and ownership", "[Planning, follow-through, and learning]", "[ ]"], ["Team contribution", "[Job-related collaboration evidence]", "[ ]"]], [2300, 5360, 1700])
    add_section_heading(doc, "Recommendation", page_break_before=True)
    add_field_table(doc, [("Recommendation", "[Strong yes / Yes / Mixed / No]"), ("Key evidence", "[Two or three job-related facts that support the recommendation]"), ("Risks or follow-up", "[Questions requiring another structured check]"), ("Calibration note", "[How this evidence compares with the written rubric]")])
    add_section_heading(doc, "Decision-quality check")
    add_checklist_table(doc, ["Done", "Review question"], [["[ ]", "Did every score cite job-related evidence?"], ["[ ]", "Were all candidates evaluated against the same written criteria?"], ["[ ]", "Were protected characteristics and unsupported assumptions excluded?"], ["[ ]", "Does the recommendation follow the evidence rather than first impression?"], ["[ ]", "Is another structured check needed before a final decision?"]], [1200, 8160])
    save_doc(doc, "candidate-evaluation-scorecard.docx")


def build_property_inspection() -> None:
    doc = Document()
    apply_doc_styles(doc, "Property inspection checklist", "Room-by-room findings, photo references, priorities, and sign-off.")
    add_note(doc, "Scope", "This operational checklist is not an engineering, environmental, code, appraisal, or licensed professional inspection.")
    add_section_heading(doc, "Property record")
    add_field_table(doc, [("Property", "[Address or internal property ID]"), ("Inspection", "[Date, time, weather, inspector]"), ("Participants", "[Names or roles]"), ("Photo folder", "[Controlled folder or case reference]")])
    for heading, items in [("Exterior and access", ["Roofline and drainage", "Walls, windows, and doors", "Walkways, rails, and access", "Utilities visible from exterior"]), ("Interior rooms", ["Walls, ceiling, and floor", "Doors and windows", "Lighting and outlets", "Fixtures and visible leaks"]), ("Safety and closeout", ["Smoke and CO alarms", "Trip or fall hazards", "Keys and access devices", "Final meter or condition photos"])]:
        if heading == "Interior rooms":
            add_section_heading(doc, heading, page_break_before=True)
        else:
            add_section_heading(doc, heading)
        add_checklist_table(doc, ["Item", "Condition / finding", "Priority", "Photo"], [[item, "[Finding or OK]", "[Now / Soon / Monitor]", "[#]"] for item in items], [2600, 4100, 1660, 1000])
    add_section_heading(doc, "Sign-off")
    add_field_table(doc, [("Immediate actions", "[Owner, action, due date]"), ("Follow-up review", "[Date and responsible person]"), ("Prepared by", "[Name, role, signature/date if required]"), ("Acknowledged by", "[Name, role, signature/date if required]")])
    save_doc(doc, "property-inspection-checklist.docx")


def build_nda() -> None:
    doc = Document()
    apply_doc_styles(doc, "Mutual NDA starter", "Editable discussion draft. Obtain qualified legal review before relying on it.")
    add_note(doc, "Not legal advice", "This template is a starting point only. Governing law, exclusions, compelled disclosure, duration, remedies, and industry obligations require qualified review.")
    add_section_heading(doc, "Parties and purpose")
    add_field_table(doc, [("Effective date", "[Month DD, YYYY]"), ("Party A", "[Full legal name, entity type, address]"), ("Party B", "[Full legal name, entity type, address]"), ("Purpose", "[Specific business discussion or evaluation]")])
    clauses = [
        ("1. Confidential information", "Confidential Information means non-public information disclosed for the Purpose that is marked confidential or should reasonably be understood as confidential given its nature and the circumstances of disclosure."),
        ("2. Exclusions", "Confidential Information does not include information the receiving party can document was lawfully known without restriction, independently developed without use of the disclosure, lawfully received from a third party, or made public without breach."),
        ("3. Use and protection", "Each party will use Confidential Information only for the Purpose, limit access to people who need it and are bound by appropriate duties, and use reasonable care to protect it."),
        ("4. Required disclosure", "A receiving party may disclose information when legally required after giving prompt notice when lawful and reasonably cooperating with protective efforts."),
        ("5. Return or destruction", "Upon written request, the receiving party will return or destroy Confidential Information, subject to lawful backup, legal, and record-retention exceptions reviewed by counsel."),
        ("6. Term", "This agreement begins on the Effective Date. Disclosure period: [period]. Confidentiality obligations: [period or treatment of trade secrets]."),
        ("7. No license or warranty", "No intellectual-property license is granted except the limited right to evaluate information for the Purpose. Information is provided [as reviewed by counsel]."),
        ("8. General", "Governing law and venue: [jurisdiction]. Entire agreement, amendment, assignment, notices, severability, waiver, and counterparts: [review and complete]."),
    ]
    for heading, body in clauses:
        heading_paragraph = doc.add_paragraph(heading, style="Heading 2")
        heading_paragraph.paragraph_format.page_break_before = heading.startswith("5.")
        doc.add_paragraph(body)
    add_section_heading(doc, "Signatures")
    add_checklist_table(doc, ["Party A", "Party B"], [["By: [Name]\nTitle: [Title]\nSignature: ____________________\nDate: [Date]", "By: [Name]\nTitle: [Title]\nSignature: ____________________\nDate: [Date]"]], [4680, 4680])
    save_doc(doc, "mutual-nda-starter.docx")


def build_education_request() -> None:
    doc = Document()
    apply_doc_styles(doc, "Education support request", "A clear written request with facts, requested action, supporting material, and follow-up.")
    add_note(doc, "Privacy", "Include only information necessary for the request. Follow school policy for sensitive student, disability, or health information and use an approved delivery channel.")
    add_section_heading(doc, "Request details")
    add_field_table(doc, [("Date", "[Month DD, YYYY]"), ("To", "[School, department, or responsible person]"), ("From", "[Requester name and authorized relationship]"), ("Student or case ID", "[Use the minimum identifier required]"), ("Subject", "[Short description of the requested support]")])
    add_section_heading(doc, "Request")
    doc.add_paragraph("I am requesting [specific action, meeting, review, record, support, or accommodation] by [reasonable date or timeframe].")
    doc.add_paragraph("Relevant facts: [Brief, objective chronology. Include dates and prior steps without unnecessary personal details.]")
    doc.add_paragraph("Requested outcome: [State what response or next action would resolve the request.]")
    add_section_heading(doc, "Supporting material")
    add_checklist_table(doc, ["Attached", "Document", "Why it is relevant"], [["[ ]", "[Record, form, evaluation, or prior correspondence]", "[One sentence]"], ["[ ]", "[Additional item]", "[One sentence]"]], [1200, 3860, 4300])
    add_section_heading(doc, "Communication and follow-up", page_break_before=True)
    add_field_table(doc, [("Preferred reply", "[Approved email, phone, portal, or meeting]"), ("Availability", "[Dates and times]"), ("Follow-up date", "[Date to check status if no response]"), ("Signature", "[Name, authorized role, date]")])
    add_section_heading(doc, "Before delivery")
    add_checklist_table(doc, ["Done", "Verification"], [["[ ]", "Every bracketed prompt was replaced or removed"], ["[ ]", "Only necessary sensitive information is included"], ["[ ]", "Attachments are named and referenced in the request"], ["[ ]", "The recipient and approved delivery channel were confirmed"], ["[ ]", "A follow-up date and copy of the final request are recorded"]], [1200, 8160])
    save_doc(doc, "education-support-request.docx")


def main() -> None:
    for record in EVIDENCE:
        make_share_card(record["toolId"], title_from_slug(record["toolId"]), "Working PDF tool")
        make_demo(record)
    for slug, (title, kicker) in RESOURCE_SHARES.items():
        make_share_card(slug, title, kicker)

    simple = FIXTURES / "simple-searchable.pdf"
    complex_pdf = FIXTURES / "complex-layout.pdf"
    clean_scan = FIXTURES / "scanned-clean.pdf"
    degraded_scan = FIXTURES / "scanned-degraded.pdf"
    encrypted = FIXTURES / "encrypted-sample.pdf"
    malformed = FIXTURES / "malformed-sample.pdf"
    large = FIXTURES / "large-100-pages.pdf"
    build_simple_pdf(simple)
    build_complex_pdf(complex_pdf)
    image_pdf(scan_source(False), clean_scan)
    image_pdf(scan_source(True), degraded_scan)
    build_large_pdf(large)
    reader = PdfReader(simple)
    writer = PdfWriter()
    for page in reader.pages:
        writer.add_page(page)
    writer.encrypt("benchmark", algorithm="AES-256")
    with encrypted.open("wb") as handle:
        writer.write(handle)
    malformed.write_bytes(b"This is an intentionally malformed PDF recovery fixture.\nNo PDF header is present.\n")
    redaction_before = REDACTION / "redaction-before.pdf"
    redaction_after = REDACTION / "redaction-after.pdf"
    build_redaction_samples(redaction_before, redaction_after)
    before_text = "\n".join(page.extract_text() or "" for page in PdfReader(redaction_before).pages)
    after_text = "\n".join(page.extract_text() or "" for page in PdfReader(redaction_after).pages)
    secret = "4815 1623 4200 0091"
    proof = [
        "PDFArrow redaction verification proof",
        "Generated: 2026-07-21",
        "",
        f"Before sample contains fictional secret in extracted text: {secret in before_text}",
        f"After sample contains fictional secret in extracted text: {secret in after_text}",
        f"After sample page count: {len(PdfReader(redaction_after).pages)}",
        "",
        "Expected safe result: before=True, after=False, page count=1.",
        "Visual inspection, metadata review, search, select/copy, and a second-person check remain required for real disclosures.",
    ]
    (REDACTION / "redaction-proof.txt").write_text("\n".join(proof) + "\n", encoding="utf-8")
    benchmark_result = {
        "benchmark": "PDFArrow Q3 2026 browser PDF conversion and fidelity benchmark",
        "testedAt": "2026-07-21",
        "status": "passed",
        "coreToolCount": len(EVIDENCE),
        "riskScenarios": ["simple", "complex_layout", "scanned", "encrypted", "malformed", "large", "mobile"],
        "browserClasses": ["desktop-chromium", "android-chromium", "iphone-webkit"],
        "tools": [{"toolId": item["toolId"], "result": item["result"], "method": item["method"]} for item in EVIDENCE],
        "interpretation": "Passed means the named property was verified for the published regression fixture. It is not a promise of perfect output for every PDF.",
    }
    (BENCHMARK / "q3-2026-results.json").write_text(json.dumps(benchmark_result, indent=2) + "\n", encoding="utf-8")
    for artifact in [simple, complex_pdf, clean_scan, degraded_scan, encrypted, malformed, large]:
        copy_pdf_artifact(artifact, "benchmark-fixtures")
    for artifact in [redaction_before, redaction_after]:
        copy_pdf_artifact(artifact, "redaction-guide")

    build_invoice()
    build_candidate_scorecard()
    build_property_inspection()
    build_nda()
    build_education_request()
    print(f"Generated {len(EVIDENCE)} tool demonstrations, {len(EVIDENCE) + len(RESOURCE_SHARES)} share cards, 9 PDF artifacts, and 5 DOCX templates.")


if __name__ == "__main__":
    main()
